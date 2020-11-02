# coding:utf-8
#!/user/bin/python3

import os.path
import os
import docx
from docx import Document
import sys
import mypatent as mp
import asyncio
import aiomysql

WORD_MIN_COUNTS = 5 
PHRASE_MIN_COUNTS = 2

count_dict = {}
count_dict_re = {}
count_multi_words = {}
count_keyword = {}

DEBUG = True
omit_list =[]
omit_words_list = []
keyword_dict = {}

def ClearDict():
    global count_dict
    global count_dict_re
    global count_multi_words
    global count_keyword

    count_dict.clear()
    count_dict_re.clear()
    count_multi_words.clear()
    count_keyword.clear()

def Debug(*args):
    #DEBUG = False
    global DEBUG
    if DEBUG == False:
        return

    s = ""
    for i in args:
        if type(i) == str:
            s= s+ i
        else:
            s = s + str(i)
    print(s)

async def prepareDict( db ):
    global omit_list
    global omit_words_list
    global keyword_dict

    rs = await mp.get_all_omit_words( db)
    for r in rs:
        omit_list.append(r.word)
    rs = await mp.get_all_omit_phrases( db )
    for r in rs:
        omit_words_list.append( r.phrase )
    rs = await mp.get_all_keywords( db )
    for r in rs:
        keyword_dict[r.word] = r.id

def KeywordFind(str):
    global count_keyword
    for k in keyword_dict.keys():
        fc = str.count(k)
        if  fc > 0:
            if k in count_keyword.keys():
                count_keyword[ k ] = count_keyword[ k ] + fc 
            else:
                count_keyword[ k ] = fc

def WordCount(str):
    # 文章字符串前期处理
    global count_dict
    strl_ist = str.replace('\n', ' ').lower().split(' ')
    # 如果字典里有该单词则加1，否则添加入字典
    for str in strl_ist:
        if str in count_dict.keys():
            count_dict[str] = count_dict[str] + 1
        else:
            count_dict[str] = 1
    #按照词频从高到低排列

def WordsLink(strs, pref=''):
    # 连续idx的单词组成短语
    global count_dict_re
    global count_multi_words

    idx_link_list = {}  
    #   index -> word
    strl_ist = strs.replace('\n', ' ').lower().split(' ')
    max_idx = len(strl_ist)
    for i in range( max_idx ):
        if strl_ist[i] in count_dict_re.keys():
            k = str(i+1)
            Debug(pref, "; idx=", i+1 , " ：" ,strl_ist[i], ", counts=", count_dict_re[strl_ist[i]])
            idx_link_list[k] = strl_ist[i]

    mulWords=""
    Debug(idx_link_list)

    intab = "!,.;:?})][("
    outab = "           "
    trantab = str.maketrans(intab, outab)

    for j in range( max_idx ):
        if str(j+1) in idx_link_list.keys():
            word = idx_link_list[str(j+1)]
            new_word = word.translate(trantab)
            new_word = new_word.strip()
            Debug("new_word = ", new_word, ";  word = ", word)

            if len(new_word) > 0:
                if len(mulWords) > 0 :
                    mulWords = mulWords + " " + new_word
                else:
                    mulWords = new_word
            if len(new_word) < len(word) and len(mulWords) > 0:
                Debug("new_word=", new_word,"; word=", word, "; Multi-worlds : [[ ----- ", mulWords, " ----- ]]; length=", len(mulWords))
                if mulWords.count(' ') > 0:
                    if mulWords in count_multi_words.keys():
                        count_multi_words[mulWords] = count_multi_words[mulWords] + 1
                    else:
                        count_multi_words[mulWords] = 1
                mulWords = ""
        else:
            if len(mulWords) > 0 and  mulWords.count(' ') > 0 :
                Debug("word [", word,"] is not in result key words, stop!!! ")
                Debug("Multi-words : [[ -----  ", mulWords ," -----  ]]") 
                if mulWords in count_multi_words.keys():
                    count_multi_words[mulWords] = count_multi_words[mulWords] + 1
                else:
                    count_multi_words[mulWords] = 1
            mulWords = ""
            continue

    # The end!!
    if len(mulWords) > 0 and mulWords.count(' ') > 0:
        Debug("count(' ')>0, so  Multi-worlds : [[ -----  ", mulWords, "  ----- ]], length = ", len(mulWords))
        if mulWords in count_multi_words.keys():
            count_multi_words[mulWords] = count_multi_words[mulWords] + 1 
        else: 
            count_multi_words[mulWords] = 1
    else:
        Debug("count(' ') ==0 or len(mulWords) = 0 : mulWords = ", mulWords)


def WordFilter():
    global count_dict
    global count_dict_re
    global WORD_MIN_COUNTS


    intab = "!,.;:?})]([{"
    outab = "            "
    trantab = str.maketrans(intab, outab)


    keys = count_dict.keys()
    values = count_dict.values()

    list_one = [(key, val) for key, val in zip(keys, values)]
    list_sort = sorted(list_one, key=lambda x: x[1], reverse=True)
    count_dict_re.clear()
    for one in list_sort:
        if int(one[1]) < WORD_MIN_COUNTS :
            pass
        else:
            new_word = one[0]
            new_word = new_word.translate(trantab)
            new_word = new_word.strip()

            if len(new_word) > 1:
                count_dict_re[new_word] = one[1]
   #按照词频从高到低排列

async def OneWordSave2Db( db, docId ):
    global count_dict_re
    global omit_list

    keys = count_dict_re.keys()
    for k in keys:
        if k in omit_list:
            continue
        else:
            await mp.add_doc_word(db, docId, k, count_dict_re[k])
                

async def PhraseSave2Db( db, docId ):
    global PHRASE_MIN_COUNTS 
    global count_multi_words
    global omit_words_list
    keys = count_multi_words.keys()
    values = count_multi_words.values()

    list_one = [(key, val) for key, val in zip(keys, values)]
    list_sort = sorted(list_one, key=lambda x: x[1], reverse=True)

    result = {}
    for one in list_sort:
        if one[0] in omit_words_list:
            pass
        elif  int(one[1]) >= PHRASE_MIN_COUNTS:
            await mp.add_doc_phrase(db, docId, one[0], one[1])

async def DocKeywordSave2Db( db, docId ):
    global count_keyword
    global keyword_dict

    keys = count_keyword.keys()
    for k in keys:
        if k in keyword_dict.keys():
            kwId = keyword_dict[ k ]
            await mp.add_doc_kw_word(db, k, count_keyword[k], docId, kwId )
        else:
            continue


def OneWord2Out(docId ):
    global count_dict_re
    global omit_list

    keys = count_dict_re.keys()
    for k in keys:
        if k in omit_list:
            continue
        else:
            print("docId=",docId,"; word=", k, "; counts=", count_dict_re[k])
                
def Keyword2Out(docId ):
    global count_keyword
    global keyword_dict

    keys = count_keyword.keys()
    for k in keys:
        if k in keyword_dict.keys():
            kwId = keyword_dict[ k ]
            print("docId=",docId,"; kwId = ",kwId, "; keyword=", k, "; counts=", count_keyword[k])
        else:
            continue
 
def Phrase2Out( docId ):
    global PHRASE_MIN_COUNTS 
    global count_multi_words
    global omit_words_list
    keys = count_multi_words.keys()
    values = count_multi_words.values()

    list_one = [(key, val) for key, val in zip(keys, values)]
    list_sort = sorted(list_one, key=lambda x: x[1], reverse=True)

    result = {}
    for one in list_sort:
        if one[0] in omit_words_list:
            pass
        elif  int(one[1]) >= PHRASE_MIN_COUNTS:
            print("docId = ",docId,"; phrase=", one[0],"; counts = ", one[1])

async def Document_word_dig(db,  docx_file):
    document = Document( docx_file )
    f_name = os.path.basename( docx_file )
    docId = await mp.add_document(db, f_name )
    if docId <=0 :
        print("error to save document Name and get docID !")
        return docId

    for paragraph in document.paragraphs:
        KeywordFind( paragraph.text)
        WordCount(paragraph.text)
    return docId

def Document_phrase_dig( docxfile ):            
    document = Document( docxfile )
    pIdx = 1
    for paragraph in document.paragraphs:
        pref = "File:" + docxfile + " ; pragraph idx=" + str(pIdx) 
        WordsLink(paragraph.text, pref)
        pIdx +=1

async def Process( db, doc_file ):
    global DEBUG
    DEBUG = False
    ClearDict()
    await prepareDict(db)
    docId = await Document_word_dig(db,  doc_file)
    if docId > 0:
        WordFilter() 
        count_dict.clear()
        Document_phrase_dig(doc_file)
        await PhraseSave2Db( db, docId )
        await OneWordSave2Db(db, docId )
        await DocKeywordSave2Db( db, docId)
    else:
        print("Error to get docID( save db error)!")
    ClearDict() 

async def main( argv ):

    if len(argv) < 2 :
        print("python3 ",argv[0]," docxfile -d/-t/DB | python3 ", argv[0]," docx_path -D/-t/DB")
        return

    global DEBUG
    g_db = await aiomysql.connect(
        host="127.0.0.1",
        port=3306,
        user="lengsh",
        password="123456",
        db="patent",
        charset="utf8mb4",
    )

    ClearDict()
    await prepareDict(g_db)
    print(omit_words_list)
    print(omit_list)
    print(keyword_dict)

    DEBUG = False
    out_format = '-t'
    if len(argv) > 2:
        out_format = argv[2]
    if out_format == '-d' or out_format == '-D':
        DEBUG = True

    fargv = argv[1]
    if os.path.isfile( fargv ):
        if fargv.find('.docx') > 0 or fargv.find('.DOCX') > 0 :
            if out_format == 'DB':
                await Process(g_db, fargv)
            else:
                docId = await Document_word_dig(g_db, fargv)
                if docId > 0:
                    WordFilter() 
                    count_dict.clear()
                    Document_phrase_dig(fargv)
############################################################
                    Phrase2Out(docId)
                    OneWord2Out(docId)
                    Keyword2Out(docId)
                else:
                    print("Error to get docID( save db error)!")
        else:
            print("Error, not docx file!")


    if os.path.isdir(fargv):
        dirs = os.listdir(fargv)
        for docxf in dirs:
            fname = os.path.join(fargv, docxf)
            if os.path.isfile(fname) and (docxf.find('.docx') > 0 or docxf.find('.DOCX') > 0) :
                if out_format == 'DB':
                    print("Process ", docxf)
                    await Process(g_db, fname)
                else:
                    print(out_format)
                    docId = await Document_word_dig(g_db, docxf)
                    if docId > 0 :
                        WordFilter() 
                        count_dict.clear()
                        Document_phrase_dig(docxf)
                        Phrase2Out(docId)
                        OneWord2Out(docId)
                        Keyword2Out(docId)
                        ClearDict()
                    else:
                        print("Error to get docID (save db error)!")
            else:
                print('ERRORRRRR')
    g_db.close()

###################################  MAIN

if __name__ == "__main__":
    # execute only if run as a script
    asyncio.run(main( sys.argv))

